import streamlit as st
from docx import Document
import io

st.title('Patient Safety Breifing')


area_choice = st.selectbox('AREA', ['CHUB', 'GREATER MANCHESTER', 'LANCASHIRE'])

if area_choice == 'CHUB':
    pb_choice = st.radio('PB', ['Fine', 'No Contact'])
    if pb_choice == 'No Contact':
        # st.caption('PB Info')
        st.text_area('PB Info')


# st.write('Thankyou for picking {} '.format(area_choice))

if st.button('save'):
    document = Document()

    document.add_heading('Thankyou for picking {} '.format(area_choice), 0)

    document.save('demo.docx')

# st.download_button() 'Download Word doc', 'demo.docx')


    # doc_download = doc_file_creation(doc_file)

    bio = io.BytesIO()
    document.save(bio)
    if document:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name="Report.docx",
            mime="docx"
        )
    # with open('demo.docx', 'rb') as pdf_file:
    #     PDFbyte = Document(pdf_file)

    # st.download_button('Download Word doc', data = PDFbyte, file_name = 'psb.docx', use_container_width=True, mime='docx')